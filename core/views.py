import json
import os

from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.views.decorators.http import require_POST

from docx import Document
from docx.shared import Inches

from .models import MeetingInput, Frame, Runbook
from .utils import (
    extract_frames,
    extract_audio,
    transcribe,
    extract_pdf_text,
    remove_duplicate_frames,
    smart_select_frames,
    generate_runbook_ai,
    normalize_template_with_gemini,
    answer_question_ai,
    edit_runbook_ai,
)


# ---------------------------------------------------------------
# UPLOAD + GENERATE
# ---------------------------------------------------------------

def upload(request):
    if request.method == "POST":
        # ---- INPUTS ----
        video = request.FILES.get("video")
        pdf   = request.FILES.get("pdf")
        template_text = request.POST.get("template", "")

        # ---- CREATE MEETING ENTRY ----
        meeting = MeetingInput.objects.create(
            video=video,
            pdf=pdf,
            status="uploaded",
        )

        # ---- TEMPLATE SCHEMA ----
        if template_text.strip():
            template_schema = normalize_template_with_gemini(template_text)
        else:
            template_schema = {
                "sections": [
                    {"heading": "Overview"},
                    {"heading": "Discussion"},
                    {"heading": "Key Decisions"},
                    {"heading": "Action Items"},
                    {"heading": "Risks"},
                    {"heading": "Next Steps"},
                ]
            }

        meeting.template_schema = template_schema
        meeting.save()

        # ---- PDF TEXT EXTRACTION ----
        pdf_text = ""
        if meeting.pdf:
            pdf_text = extract_pdf_text(meeting.pdf.path)

        # ---- VIDEO → FRAMES ----
        extract_frames(meeting.video.path, meeting)
        remove_duplicate_frames("media/frames")
        selected_frames = smart_select_frames("media/frames")

        meeting.selected_frames = selected_frames
        meeting.save()
        print("SMART FRAMES:", selected_frames)

        # ---- VIDEO → AUDIO → TRANSCRIPT ----
        audio_path          = extract_audio(meeting.video.path)
        transcription_result = transcribe(audio_path)

        transcript        = transcription_result["transcript"]
        detected_language = transcription_result["detected_language"]
        was_translated    = transcription_result["was_translated"]

        if was_translated:
            print(f"TRANSLATION: detected '{detected_language}', translated to English")
        print("TRANSCRIPT:", transcript[:200])

        meeting.transcript        = transcript
        meeting.detected_language = detected_language
        meeting.was_translated    = was_translated
        meeting.status            = "processed"
        meeting.save()

        # ---- SAVE TRANSCRIPT AS .TXT ----
        txt_path = f"media/transcript_{meeting.id}.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"Meeting ID: {meeting.id}\n")
            f.write(f"Detected language: {detected_language.upper()}")
            if was_translated:
                f.write(" (translated to English)")
            f.write("\n\n")
            f.write(transcript)

        # ---- COMBINE INPUTS FOR AI ----
        combined_text = f"Transcript:\n{transcript}\n\nPDF Content:\n{pdf_text}"

        # ---- AI RUNBOOK GENERATION ----
        runbook_data = generate_runbook_ai(
            combined_text, meeting.selected_frames, template_schema
        )

        # ---- BUILD DOCX ----
        _build_docx(meeting.id, runbook_data)

        Runbook.objects.create(meeting=meeting, content=runbook_data)

        return redirect("success", meeting_id=meeting.id)

    return render(request, "upload.html")


# ---------------------------------------------------------------
# FRAME PREVIEW
# ---------------------------------------------------------------

def frame_preview(request, meeting_id):
    meeting = get_object_or_404(MeetingInput, id=meeting_id)
    frames  = Frame.objects.filter(
        meeting=meeting, image__in=meeting.selected_frames
    )
    return render(request, "frames.html", {"frames": frames})


# ---------------------------------------------------------------
# SUCCESS PAGE
# ---------------------------------------------------------------

def success(request, meeting_id):
    meeting = get_object_or_404(MeetingInput, id=meeting_id)
    runbook = get_object_or_404(Runbook, meeting=meeting)
    return render(request, "success.html", {
        "meeting": meeting,
        "runbook": runbook.content,
    })


# ---------------------------------------------------------------
# TRANSCRIPT DOWNLOAD
# ---------------------------------------------------------------

def download_transcript(request, meeting_id):
    meeting  = get_object_or_404(MeetingInput, id=meeting_id)
    txt_path = f"media/transcript_{meeting_id}.txt"

    if not os.path.exists(txt_path):
        # Regenerate if missing
        with open(txt_path, "w", encoding="utf-8") as f:
            lang = (meeting.detected_language or "unknown").upper()
            f.write(f"Meeting ID: {meeting.id}\n")
            f.write(f"Detected language: {lang}")
            if meeting.was_translated:
                f.write(" (translated to English)")
            f.write("\n\n")
            f.write(meeting.transcript or "")

    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    response = HttpResponse(content, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = (
        f'attachment; filename="transcript_{meeting_id}.txt"'
    )
    return response


# ---------------------------------------------------------------
# EDIT RUNBOOK  (GET = show AI editor, POST = apply instruction via Gemini)
# ---------------------------------------------------------------

def edit_runbook(request, meeting_id):
    meeting = get_object_or_404(MeetingInput, id=meeting_id)
    runbook = get_object_or_404(Runbook, meeting=meeting)

    if request.method == "POST":
        try:
            payload = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON"}, status=400)

        instruction = payload.get("instruction", "").strip()
        if not instruction:
            return JsonResponse({"error": "instruction is required"}, status=400)

        # Use the current_runbook from payload if provided (client keeps latest state),
        # otherwise fall back to what's stored in the DB.
        current_runbook = payload.get("current_runbook") or runbook.content

        try:
            updated_runbook = edit_runbook_ai(instruction, current_runbook)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

        # Persist updated runbook to DB
        runbook.content = updated_runbook
        runbook.save()

        # Rebuild .docx with the updated content
        _build_docx(meeting_id, updated_runbook)

        return JsonResponse({"ok": True, "runbook": updated_runbook})

    # GET — render the AI editor
    return render(request, "edit_runbook.html", {
        "meeting": meeting,
        "runbook": runbook.content,
        "runbook_json": json.dumps(runbook.content),
    })


# ---------------------------------------------------------------
# Q&A  (POST only — returns JSON answer)
# ---------------------------------------------------------------

@require_POST
def qa(request, meeting_id):
    meeting = get_object_or_404(MeetingInput, id=meeting_id)

    try:
        payload  = json.loads(request.body)
        question = payload.get("question", "").strip()
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    if not question:
        return JsonResponse({"error": "Question is required"}, status=400)

    # Build context from transcript + PDF text
    transcript = meeting.transcript or ""
    pdf_text   = ""
    if meeting.pdf:
        try:
            from .utils import extract_pdf_text
            pdf_text = extract_pdf_text(meeting.pdf.path)
        except Exception:
            pass

    context = f"Transcript:\n{transcript}\n\nPDF Content:\n{pdf_text}"

    try:
        answer = answer_question_ai(question, context)
        return JsonResponse({"answer": answer})
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


def qa_page(request, meeting_id):
    meeting = get_object_or_404(MeetingInput, id=meeting_id)
    return render(request, "qa.html", {"meeting": meeting})


# ---------------------------------------------------------------
# INTERNAL HELPER — build / rebuild .docx
# ---------------------------------------------------------------

def _build_docx(meeting_id, runbook_data):
    doc = Document()
    doc.add_heading(runbook_data.get("title", "Meeting Runbook"), level=1)

    for section in runbook_data.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=2)
        for p in section.get("paragraphs", []):
            doc.add_paragraph(p)
        for img in section.get("supporting_images", []):
            img_path = os.path.join("media", img.get("image", ""))
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph(img.get("caption", ""))

    if runbook_data.get("key_decisions"):
        doc.add_heading("Key Decisions", level=2)
        for d in runbook_data["key_decisions"]:
            doc.add_paragraph(d, style="List Bullet")

    if runbook_data.get("action_items"):
        doc.add_heading("Action Items", level=2)
        for a in runbook_data["action_items"]:
            doc.add_paragraph(
                f"{a.get('task', '')} — {a.get('owner', '')} — {a.get('deadline', '')}"
            )

    if runbook_data.get("risks_or_blockers"):
        doc.add_heading("Risks / Blockers", level=2)
        for r in runbook_data["risks_or_blockers"]:
            doc.add_paragraph(r, style="List Bullet")

    if runbook_data.get("next_steps"):
        doc.add_heading("Next Steps", level=2)
        for n in runbook_data["next_steps"]:
            doc.add_paragraph(n, style="List Bullet")

    doc.save(f"media/runbook_{meeting_id}.docx")